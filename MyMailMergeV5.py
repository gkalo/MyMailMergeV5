import re
import io
import logging
import argparse
import traceback
import sys
import os
import datetime
import smtplib
import getpass

from inspect import getframeinfo, stack

import yaml

import pyexcel
#if no oracle connections are needed, comment out next 1 line(s)
import cx_Oracle
#if no mysql connections are needed, comment out next 2 line(s)
import mysql.connector as mysql
from mysql.connector import FieldType as mysql_FieldType

from mailmerge import MailMerge
from docx import Document
from docx.shared import Mm

from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from email.utils import formataddr
from email.utils import formatdate
from email.header import Header
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage

wFileName = 'MyMailMergeV5'
    
argParser = argparse.ArgumentParser()
argParser.add_argument('--config', '-c', help="Configuration file to be used. Default = '" + wFileName + ".yml'", type = str, default = '-')
argParser.add_argument('--input',  '-i', help="Input file (template file) to be used. Default = '" + wFileName + ".docx'", type = str, default = '-')
argParser.add_argument('--output', '-o', help="Output file to be created. Default = '" + wFileName + "_Out.docx'", type = str, default = '-')
argParser.add_argument('--log',    '-l', help="Logging file to be used. Default = ''", type = str, default = '-')
argParser.add_argument('--mailTemplate_html', '-mh', help="Mail Template HTML file to be used. Default = ''", type = str, default = '-')
argParser.add_argument('--mailTemplate_text', '-mt', help="Mail Template Text file to be used. Default = ''", type = str, default = '-')
argParser.add_argument('--attachments', '-a', help="List of Attachments for Mail. Default = []", type = list, default = ['-'])
argParser.add_argument('--images_attached', '-im', help="List of Images attached to HTML template. Default = []", type = list, default = ['-'])

args = argParser.parse_args()

#----------------------------------------------------------------------------------------------------------------------
#examine command line arguments and configuration file, decide what to use
def examineAllArgs():
    global cfg, mmm3_query, recs_cnt
    
    if args.config == argParser.get_default('config'):
        args.config = wFileName + '.yml'
    with open(args.config, "r", encoding="utf8") as ymlfile:
        cfg = yaml.load(ymlfile, Loader=yaml.SafeLoader)
    
    mmm3_query = list(cfg["Queries"])
    for i in range(len(mmm3_query)):
        recs_cnt.append(0)
    
    if args.input == argParser.get_default('input'):
        if 'input' in cfg["Params"]:
            args.input = cfg["Params"]['input']
        else:
            args.input = wFileName + '.docx'
    if args.output == argParser.get_default('output'):
        if 'output' in cfg["Params"]:
            args.output = cfg["Params"]['output']
        else:
            args.output = wFileName + '_Out.docx'
    if args.log == argParser.get_default('log'):
        if 'log' in cfg["Params"]:
            args.log = cfg["Params"]['log']
        else:
            args.log = ''
    if args.mailTemplate_html == argParser.get_default('mailTemplate_html'):
        if 'mailTemplate_html' in cfg["Params"]:
            args.mailTemplate_html = cfg["Params"]['mailTemplate_html']
        else:
            args.mailTemplate_html = ''
    if args.mailTemplate_text == argParser.get_default('mailTemplate_text'):
        if 'mailTemplate_text' in cfg["Params"]:
            args.mailTemplate_text = cfg["Params"]['mailTemplate_text']
        else:
            args.mailTemplate_text = ''
    if args.attachments == argParser.get_default('attachments'):
        if 'attachments' in cfg["Params"]:
            args.attachments = cfg["Params"]['attachments']
        else:
            args.attachments = []
    if args.images_attached == argParser.get_default('images_attached'):
        if 'images_attached' in cfg["Params"]:
            args.images_attached = cfg["Params"]['images_attached']
        else:
            args.images_attached = []
    
    myLogging("Open", "")

    myLogging("Debug", "Arguments")
    myLogging("Debug", "    config           =" + args.config)
    myLogging("Debug", "    input            =" + args.input)
    myLogging("Debug", "    output           =" + args.output)
    myLogging("Debug", "    log              =" + args.log)
    myLogging("Debug", "    mailTemplate_html=" + args.mailTemplate_html)
    myLogging("Debug", "    mailTemplate_text=" + args.mailTemplate_text)
    myLogging("Debug", "    attachments      =" + str(args.attachments))
    myLogging("Debug", "    images_attached  =" + str(args.images_attached))
    if 'Images' in cfg:
        myLogging("Debug", "    Images=           " + str(cfg["Images"]))
#----------------------------------------------------------------------------------------------------------------------
#puts in log the lMsg with level equal to lLevel
def myLogging(lLevel, lMsg):
    global log_file
    if args.log:
        caller = getframeinfo(stack()[1][0])
        wMsg = "%s:%s:%s:%s:%s" % (datetime.datetime.now(), lLevel, stack()[1][3].ljust(25), str(caller.lineno).rjust(3), lMsg)
        if lLevel.lower() == 'open':
            log_file = open(args.log,'w')
        elif lLevel.lower() == 'debug':
            log_file.write(wMsg+'\n')
        elif lLevel.lower() == 'info':
            log_file.write(wMsg+'\n')
            print("Info:" + str(lMsg))
        elif lLevel.lower() == 'error':
            log_file.write(wMsg+'\n')
            print("Error:" + str(lMsg))
        else:
            print("Invalid logging Level = '" + lLevel + "'")
#----------------------------------------------------------------------------------------------------------------------
#every row is converted to dictionary (key:value pair). col[0] is column name in cursor description.
#column name must be lower, lower must also be the names of merge fields in word.
# Case of column names depends on settings in DB, so convert everything to lower to be safe.
def RowToDict(q_no, row, cur, q_images, q_regEx):
    qSeq = 'q%s_' % q_no
    global img_cnt

    if cur:
        rec = dict(zip([col[0].lower() for col in cur.description], row))
        wIter = cur.description
    else:
        rec=row
        wIter = rec.keys()
    if recs_cnt[q_no-1] == 1:
        myLogging("Debug", rec.keys())
    
    for col in wIter:
        isImage = False
        col0 = ""
        colCont = ""
        if cur:
#if no oracle connections are needed, comment out next 5 line(s)
            if dbType == "oracle":
                if col[1] == cx_Oracle.BLOB and rec[col[0].lower()]:
                    col0 = col[0].lower()
                    colCont = "<BLOB>"
                    isImage = True
#if no mysql connections are needed, comment out next 5 line(s)
            elif dbType == "mysql":
                if "blob" in mysql_FieldType.get_info(col[1]).lower() and rec[col[0].lower()]:
                    col0 = col[0].lower()
                    colCont = "<BLOB>"
                    isImage = True
        if not cur and 'Images' in cfg and col.lower() in cfg['Images']:
            col0 = col.lower()
            colCont = rec[col0]
            isImage = True
        if isImage:
            img_cnt += 1
            imgKey = '{{%simage_%s}}' % (qSeq, img_cnt)
            regExp = "\{\{%simage_%s\}\}" % (qSeq, img_cnt)
            myLogging("Debug", "q_no=" + str(q_no) + ", " + imgKey + " = " + colCont)
            q_images[imgKey]=rec[col0]
            q_regEx[imgKey]=regExp
            rec[col0] = imgKey
            if 'Images' in cfg and col0 in cfg['Images']:
                q_images[imgKey+"_dim"] = cfg['Images'][col0]
    
    return rec, q_images, q_regEx
#----------------------------------------------------------------------------------------------------------------------
#it will replace a regex string with an image
def docx_replace_regex(doc_obj, regex, thisImage, imageKey):
 
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            # add_picture will not work in paragraphs, not supported!
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    w=None
                    h=None
                    if imageKey+"_dim" in images:
                        dim = images[imageKey+"_dim"]
                        if 'width' in dim and dim['width']:
                            w = Mm(int(dim['width']))
                        if 'height' in dim and dim['height']:
                            h = Mm(int(dim['height']))
                    inline[i].text = regex.sub("", inline[i].text)
                    inline[i].add_picture(thisImage,width=w,height=h)
                    
 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , thisImage, imageKey)
#----------------------------------------------------------------------------------------------------------------------
#Get a value for every Bind Variable
def GetValueForBindVars(bindParams, wRec):
    bindVars = {}
    for bindVar in bindParams:
        if bindVar.lower() in wRec:
            bindVars[bindVar] = wRec[bindVar.lower()]
        else:
            bindVarVal = input("Please enter value for " + bindVar + ": ")
            bindVars[bindVar] = bindVarVal
    
    return bindVars
#----------------------------------------------------------------------------------------------------------------------
#Prepare an SQL statement and replace Bind Variables
def QueryPrepAndBind(q_no, q_cur, q_sql):
    bindVars = {}
#if no oracle connections are needed, comment out next 3 line(s)
    if dbType == "oracle":
        q_cur.prepare(q_sql)
        bindVars = GetValueForBindVars(q_cur.bindnames(), rec_mas)
#if no mysql connections are needed, comment out next 2 line(s)
    elif dbType == "mysql":
        bindVars = GetValueForBindVars(re.findall(r"\%\((.*?)\)s", q_sql), rec_mas)
        
    myLogging("Debug", "bindVars=" + str(bindVars))
    return bindVars
#----------------------------------------------------------------------------------------------------------------------
def ReplaceBindVars(wCondition):
    #check for bind variables
    for bindVar in re.findall(r"\%\((.*?)\)s", wCondition):
        if bindVar.lower() in rec_mas:
            bindVarVal = rec_mas[bindVar.lower()]
        else:
            bindVarVal = input("Please enter value for " + bindVar + ": ")
            
        wCondition = re.sub(r"\%\((.*?)\)s", bindVarVal, wCondition, flags=re.IGNORECASE)
    
    return wCondition
#----------------------------------------------------------------------------------------------------------------------
def GetSheetFromSpreadSheet(q_no):
    wFile_name = mmm3_query[q_no-1]['file_name']
    if 'sheet' in mmm3_query[q_no-1] and mmm3_query[q_no-1]['sheet']:
        wSheet_name = mmm3_query[q_no-1]['sheet']
        book = pyexcel.get_book(file_name=wFile_name)
        sheet = book[wSheet_name]
        sheet.name_columns_by_row(0)
        myLogging("Debug", "Query No " + str(q_no - 1) + " is in a file named " + wFile_name + " sheet named '" + wSheet_name + "'")
    else:
        delim = '\t'
        if 'delimiter' in mmm3_query[q_no-1] and mmm3_query[q_no-1]['delimiter']:
            delim = mmm3_query[q_no-1]['delimiter']
        myLogging("Debug", "Query No " + str(q_no - 1) + " is in a delimited file named " + wFile_name + ", delimiter is [" + delim + "]")
        
        sheet = pyexcel.get_sheet(file_name=wFile_name, name_columns_by_row=0, delimiter=delim)
    
    #Convert all columns to strings, to be on the safe site. Mailmerge demands everything to be string
    for index in sheet.column_range():
        sheet.column.format(index, str)
    
    if 'filter_condition' in mmm3_query[q_no-1] and mmm3_query[q_no-1]['filter_condition']:
        cond = mmm3_query[q_no-1]['filter_condition']
        myLogging("Debug", "         " + " has also a filter condition [" + cond + "]")
        
        #replace all bind variables, if any
        cond = ReplaceBindVars(cond)
        
        myLogging("Debug", "         " + " has also a filter condition [" + cond + "]")
        
        to_remove = []
        try:
            for row in sheet.row_range(): 
                if not eval(cond):
                    to_remove.append(row)
        except Exception as error:
            myLogging("Error", "Error evaluating filter condition [" + cond + "] :" + str(error))
        finally:
            sheet.filter(row_indices=to_remove)
    
    return list(sheet.records)

#----------------------------------------------------------------------------------------------------------------------
#examine configuration Query and return Dictionary
def QueryOneToDict(q_no, q_images, q_regEx):
    global mas_cur, mmm3_query, recs_cnt
    
    recs_cnt[q_no-1] += 1
    
    if isinstance(mmm3_query[q_no-1], list):
        if mmm3_query[q_no-1]:
            if recs_cnt[q_no-1] == 1:
                myLogging("Debug", "Query No " + str(q_no - 1) + " is a list with " + str(len(mmm3_query[q_no-1])) + " elements")
            return RowToDict(q_no, mmm3_query[q_no-1].pop(0), None, q_images, q_regEx)
    if isinstance(mmm3_query[q_no-1], dict):
        if mmm3_query[q_no-1]:
            if recs_cnt[q_no-1] == 1:
                mmm3_query[q_no-1] = GetSheetFromSpreadSheet(q_no)
                
            return RowToDict(q_no, mmm3_query[q_no-1].pop(0), None, q_images, q_regEx)
    if isinstance(mmm3_query[q_no-1], str):
        if recs_cnt[q_no-1] == 1:
            myLogging("Debug", "Query No " + str(q_no - 1) + " is a query")
            
#if no oracle connections are needed, comment out next 2 line(s)
            if dbType == "oracle":
                mas_cur = mas_connection.cursor()
#if no mysql connections are needed, comment out next 2 line(s)
            if dbType == "mysql":
                mas_cur = mas_connection.cursor(buffered=True)
                
            mas_bindVars = QueryPrepAndBind(q_no, mas_cur, mmm3_query[q_no-1])
            if mas_bindVars:
                mas_cur.execute(mmm3_query[q_no-1], mas_bindVars)
            else:
                mas_cur.execute(mmm3_query[q_no-1])
        
        mas_row = mas_cur.fetchone()
        if mas_row is None:
            return {}, q_images, q_regEx
        
        return RowToDict(q_no, mas_row, mas_cur, q_images, q_regEx)
    
    return {}, q_images, q_regEx
#----------------------------------------------------------------------------------------------------------------------
#examine configuration Query and return list of Dictionary
def QueryAllToListOfDict(q_no, q_images, q_regEx):
    global mmm3_query
    
    if isinstance(mmm3_query[q_no-1], dict):
        mmm3_query[q_no-1] = GetSheetFromSpreadSheet(q_no)
    if isinstance(mmm3_query[q_no-1], list):
        myLogging("Debug", "Query No " + str(q_no - 1) + " is a list with " + str(len(mmm3_query[q_no-1])) + " elements")
        return FetchRowsIntoListOfDict(q_no, None, mmm3_query[q_no-1], q_images, q_regEx)
    if isinstance(mmm3_query[q_no-1], str):
        myLogging("Debug", "Query No " + str(q_no - 1) + " is a query")
        return FetchRowsIntoListOfDict(q_no, mas_connection, mmm3_query[q_no-1], q_images, q_regEx)
#----------------------------------------------------------------------------------------------------------------------
#execute Query and fetch all row into Dictionary
def FetchRowsIntoListOfDict(q_no, f_connection, f_query, q_images, q_regEx):
    global recs_cnt
    
    f_cur = None
    if f_connection:
#if no oracle connections are needed, comment out next 2 line(s)
        if dbType == "oracle":
            f_cur = f_connection.cursor()
#if no mysql connections are needed, comment out next 2 line(s)
        if dbType == "mysql":
            f_cur = f_connection.cursor(buffered=True)

        det_bindVars = QueryPrepAndBind(q_no, f_cur, f_query)
        if det_bindVars:
            f_cur.execute(f_query, det_bindVars)
        else:
            f_cur.execute(f_query)
    
    f_result=[]
    recs_cnt[q_no-1] = 0
    while True:
        f_row = None
        if f_connection:
            f_row = f_cur.fetchone()
        else:
            if f_query:
                f_row = f_query.pop(0)
        
        if f_row is None:
            break
            
        recs_cnt[q_no-1] += 1
        
        f_dic, q_images, q_regEx = RowToDict(q_no, f_row, f_cur, q_images, q_regEx)
        f_result.append(f_dic)
        
    return f_result, q_images, q_regEx
#----------------------------------------------------------------------------------------------------------------------
#to Fetch LOBs as Bytes
#if no oracle connections are needed, comment out next 3 line(s)
def OutputTypeHandler(cursor, name, defaultType, size, precision, scale):
    if defaultType == cx_Oracle.BLOB:
        return cursor.var(cx_Oracle.LONG_BINARY, arraysize=cursor.arraysize)
#----------------------------------------------------------------------------------------------------------------------
def docx_delete_paragraph(paragraph):
    #remove paragraph (with the following workaround, since .delete method does not exist)
    #https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

#----------------------------------------------------------------------------------------------------------------------
def checkToDeleteLine(parentIsCell, doc_part):
    if not parentIsCell:
        if len(doc_part.text) == 0:
            docx_delete_paragraph(doc_part)
        elif len(doc_part.text) == 1 and doc_part.text == " ":
            doc_part.text = ""

def deleteLastItemInIfCond(ifConditions):
    myLogging("Debug", "              : delete last item in ifConditions list with name=[" + ifConditions[len(ifConditions)-1]["name"] + "]")
    del ifConditions[len(ifConditions)-1]
    
    wIfName = ""
    if ifConditions:
        wIfName = ifConditions[len(ifConditions)-1]["name"]
    return wIfName

#----------------------------------------------------------------------------------------------------------------------
#to evaluate and keep or remove if - fi blocks with conditional statements
def findIfCondition(doc_obj):
    global ifConditions
    wMode = "Continue"
    wIfElseFiText = r"\{\{((if)_(.*?) (.*?)|(else)_(.*?)|(fi)_(.*?))\}\}"
    
    if isinstance(doc_obj, _Document):
        parent_elm = doc_obj.element.body
        parentIsCell = False
        parentIs = "Document: "
    elif isinstance(doc_obj, _Cell):
        parent_elm = doc_obj._tc
        parentIsCell = True
        parentIs = "Table   : "
    else:
        raise ValueError("something's not right")

    wIfName = ""
    if ifConditions:
        wIfName = ifConditions[len(ifConditions)-1]["name"]
            
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            doc_part = Paragraph(child, doc_obj)
            text = doc_part.text
            startDeleteAt = 0
            
            wMatch = re.search(wIfElseFiText, text, re.IGNORECASE)
            wIfElseFiTextFound = False
            while wMatch:
                wIfElseFiTextFound = True
                wGroups = wMatch.groups()
                wName = ""
                
                if str(wGroups[0]).lower().startswith("if"):
                    wGrp = wGroups[1]
                    wName = wGroups[2]
                    wCond = wGroups[3]
                    myLogging("Debug", "If condition found: name=[" + wName + "], condition=[" + wCond + "]")
                elif str(wGroups[0]).lower().startswith("else"):
                    wGrp = wGroups[4]
                    wName = wGroups[5]
                    myLogging("Debug", "Else found: name=[" + wName + "]")
                elif str(wGroups[0]).lower().startswith("fi"):
                    wGrp = wGroups[6]
                    wName = wGroups[7]
                    myLogging("Debug", "Fi found: name=[" + wName + "]")
                else:
                    myLogging("Error", "Problem found: [" + wGroups[0].lower() + "]")
                
                #remove what we found so not to find again
                text = text[0: wMatch.start():] + text[wMatch.end()::]
                
                if wGrp.lower() == "if":
                    if wMode == "Continue":
                        wIfName = wName
                        wIfCond = wCond
                        
                        #replace all bind variables, if any
                        wIfCond = ReplaceBindVars(wIfCond)
                        
                        wEval = False
                        try:
                            wEval = eval(wIfCond)
                            ifConditions.append({"name": wIfName, "condition": wIfCond, "eval": wEval})
                            myLogging("Debug", "             L:" +str(len(ifConditions)) + "    name=[" + wIfName + "], condition=[" + wIfCond + "] = " + str(wEval))
                        except Exception:
                            myLogging("Error", "Invalid if condition: name=[" + wIfName + "], condition=[" + wIfCond + "]")
                            
                        #if condition is true, continue
                        if wEval:
                            myLogging("Debug", "              : Just remove If block - ifName=[" + wIfName + "]")
                        else:
                            wMode = "Delete"
                            startDeleteAt = wMatch.start()
                            myLogging("Debug", "              : start deleting - ifName=[" + wIfName + "]")
                    else:
                        myLogging("Debug", "              : ignore it, keep deleting")
                            
                elif wGrp.lower() == "else":
                    #if is Continue we were in an If that was true, so from else until fi delete everything
                    if wMode == "Continue":
                        if wName == wIfName:
                            wMode = "Delete"
                            startDeleteAt = wMatch.start()
                            myLogging("Debug", "              : start deleting - ifName=[" + wIfName + "]")
                        else:
                            myLogging("Error", "Invalid nested if condition: it is=[" + wName + "], should be=[" + wIfName + "]")
                    
                    #else is Delete, so we stop deleting
                    else:
                        if wName == wIfName:
                            #this is the block we are looking for, delete everything from last start till end of else block and continue
                            myLogging("Debug", "              : stop deleting - ifName=[" + wIfName + "]")
                            text = text[0: startDeleteAt:] + text[(wMatch.end() - (wMatch.end() - wMatch.start()))::]
                            wMode = "Continue"
                        else:
                            myLogging("Debug", "              : ignore it, keep deleting, it is=[" + wName + "], should be=[" + wIfName + "]")
                elif wGrp.lower() == "fi":
                    #if is Continue we were in a true condition (if or else), just delete the fi block
                    if wMode == "Continue":
                        if wName == wIfName:
                            wIfName = deleteLastItemInIfCond(ifConditions)
                        else:
                            myLogging("Error", "Invalid nested if condition: it is=[" + wName + "], should be=[" + wIfName + "]")
                    
                    #else is Delete, so we stop deleting
                    else:
                        if wName == wIfName:
                            #delete everything from last start till end of fi block and continue
                            myLogging("Debug", "              : stop deleting - ifName=[" + wIfName + "]")
                            text = text[0: startDeleteAt:] + text[(wMatch.end() - (wMatch.end() - wMatch.start()))::]
                            wMode = "Continue"
                            wIfName = deleteLastItemInIfCond(ifConditions)
                        else:
                            myLogging("Debug", "              : ignore it, keep deleting, it is=[" + wName + "], should be=[" + wIfName + "]")
                
                wMatch = re.search(wIfElseFiText, text, re.IGNORECASE)
            #end of while wMatch:
            
            if wIfElseFiTextFound:
                if wMode == "Delete":
                    myLogging("Debug", "              : delete rest of line - ifName=[" + wIfName + "]")
                    text = text[0: startDeleteAt:] + text[len(text)::]
                
                doc_part.text = text
                checkToDeleteLine(parentIsCell, doc_part)
            else:
                if wMode == "Delete":
                    docx_delete_paragraph(doc_part)
                    myLogging("Debug", "              : delete line - ifName=[" + wIfName + "]")
            
        elif isinstance(child, CT_Tbl):
            doc_part = Table(child, doc_obj)
            
            if wMode == "Delete":
                #this is the correct way to remove a table
                doc_part._element.getparent().remove(doc_part._element)
                myLogging("Debug", "               : delete table")
            else:
                for row in doc_part.rows:
                    for cell in row.cells:
                        findIfCondition(cell)
#----------------------------------------------------------------------------------------------------------------------
def send_email(mailServer, from_name, from_addr, subject, html, text, mail_data, OutFile, attachments=[], img_list=[]):
    msg_root = MIMEMultipart('mixed')
    msg_root['Date'] = formatdate(localtime=1)
    msg_root['From'] = formataddr((from_name, from_addr))
    msg_root['To'] = formataddr((mail_data["to_name"], mail_data["to_addr"]))
    msg_root['Subject'] = subject
    
    msg_related = MIMEMultipart('related')
    msg_root.attach(msg_related)

    msg_alternative = MIMEMultipart('alternative')
    msg_related.attach(msg_alternative)

    text = text.format(**mail_data)
    msg_text = MIMEText(text, 'plain')
    msg_alternative.attach(msg_text)

    #replace all new lines with <br> for the html alternative
    text = re.sub(r"\n", r"<br>", text, flags=re.IGNORECASE)
    msg_html = MIMEText(html.format(text), 'html')
    msg_alternative.attach(msg_html)

    #if there are any images in the html part
    for i, img in enumerate(img_list):
        with open(img, 'rb') as fp:
            msg_image = MIMEImage(fp.read())
            msg_image.add_header('Content-ID', '<image{}>'.format(i))
            msg_related.attach(msg_image)

    for attachment in attachments:
        if attachment.lower() == 'output':
            attachment = OutFile
            
        fname = os.path.basename(attachment)

        with open(attachment, 'rb') as f:
            msg_attach = MIMEBase('application', 'octet-stream')
            msg_attach.set_payload(f.read())
            encoders.encode_base64(msg_attach)
            msg_attach.add_header('Content-Disposition', 'attachment',
                                  filename=(Header(fname, 'utf-8').encode()))
            msg_root.attach(msg_attach)

    mailServer.send_message(msg_root)
#----------------------------------------------------------------------------------------------------------------------
#WARNING!! All merge fields must be in lower case and should not contain spaces, use underscores.
log_file = None
mas_connection = None
mailServer = None
mailMsg_text = ""
mailMsg_html = ""
mas_cur = None
mmm3_query = None
cfg = {}
recs_cnt=[]
examineAllArgs()
rec_mas={}
try:
    if  "Connection" in cfg and 'username' in cfg["Connection"] and cfg["Connection"]['username']:
        if 'password' in cfg["Connection"]:
            wPass = cfg["Connection"]['password']
        else:
            wPass = getpass.getpass("Please enter password for user [" + cfg["Connection"]['username'] + "]: ")
        dbType = "oracle"
        if 'db' in cfg["Connection"] and cfg["Connection"]['db']:
            dbType = cfg["Connection"]['db'].lower()
        myLogging("Debug", "Opening Connection (" + dbType + ")")
    #if no oracle connections are needed, comment out next 3 line(s)
        if dbType == "oracle":
            mas_connection = cx_Oracle.connect(cfg["Connection"]['username'],wPass,cfg["Connection"]['dsn'],encoding=cfg["Connection"]['encoding'])
            mas_connection.outputtypehandler = OutputTypeHandler
    #if no mysql connections are needed, comment out next 3 line(s)
        elif dbType == "mysql":
            #use_pure = True is needed here otherwise MySQL tries (sometimes) to read blobs as utf8 strings and fails...
            mas_connection = mysql.connect(host = cfg["Connection"]['host'], user = cfg["Connection"]['username'], passwd = wPass, database = cfg["Connection"]['database'], use_pure = True)
            
    if "Mailer" in cfg and 'host' in cfg["Mailer"] and cfg["Mailer"]['host']:
        myLogging("Debug", "Opening Mail Connection host = " + cfg["Mailer"]['host'] + ":" + str(cfg["Mailer"]['port']))
        mailServer = smtplib.SMTP(cfg["Mailer"]['host'], cfg["Mailer"]['port'])
        mailServer.ehlo()
        try:
            myLogging("Debug", "Trying starttls on smtplib")
            mailServer.starttls()
            mailServer.ehlo()
            myLogging("Debug", "    Looks OK")
        except smtplib.SMTPException as e:
            myLogging("Info", e)
        if 'username' in cfg["Mailer"] and cfg["Mailer"]['username']:
            myLogging("Debug", "                        user = " + cfg["Mailer"]['user'])
            if 'password' in cfg["Mailer"]:
                wMPass = cfg["Mailer"]['password']
            else:
                wMPass = getpass.getpass("Please enter password for Mail user [" + cfg["Mailer"]['username'] + "]: ")
            
            mailServer.login(cfg["Mailer"]['username'], wMPass)
            
    myLogging("Debug", "Master Loop")
    images={}
    regEx={}
    img_cnt = 0
    rec_mas, images, regEx = QueryOneToDict(1, images, regEx)
    while rec_mas:
        if mailServer:
            with open(cfg["Params"]['mailTemplate_text'], encoding="utf8") as mt:
                mailMsg_text = mt.read()
            with open(cfg["Params"]['mailTemplate_html'], encoding="utf8") as mt:
                mailMsg_html = mt.read()
            
        mmDoc = MailMerge(open(args.input, 'rb'))
        if recs_cnt[0] == 1:
            myLogging("Debug", "merge_fields = " + str(sorted(mmDoc.get_merge_fields())))
            for f in mmDoc.get_merge_fields():
                if f != f.lower():
                    print("")
                    print("'%s' is not lower case. It will not be replaced. All merge fields must be in lower case." % f)

        myLogging("Debug", "master(" + str(recs_cnt[0]) + ") = " + str(rec_mas))
        #NOTE if you want to use a dictionary, put the two stars
        mmDoc.merge(**rec_mas)

        rec_det=[]

        for q_i in range(1, len(mmm3_query)):
            wRecs, images, regEx = QueryAllToListOfDict(q_i+1, images, regEx)
            #rec_det contains a list with so many rows as the number of detail queries. Every row of that list contains a list of records returned by that query
            rec_det.append([]) #append an empty list

            #NOTE you can use any key name from dictionary i.e. trans_date, trans_debit, trans_credit, trans_balance. It is easier to use first key.
            myLogging("Debug", "rec_det[" + str(q_i - 1) + "] = " + str(wRecs))
            if wRecs:
                mmDoc.merge_rows(list(wRecs[0].keys())[0], wRecs)
                rec_det[q_i - 1] = wRecs #replace the previous appended empty list with all detail records
                
        interm_stream = io.BytesIO()
        mmDoc.write(interm_stream)
        
        #NOTE if you want to take a look in the intermediate document, un-comment next two lines
        #intDoc = Document(interm_stream)
        #intDoc.save(wFileName + '_Int.docx')
        
        dxDoc = Document(interm_stream)
        
        for r in regEx.keys():
            regex1 = re.compile(regEx[r])
            if isinstance(images[r], str):
                docx_replace_regex(dxDoc, regex1, images[r], r)
            else:
                ##NOTE this conversion plus the OutputTypeHandler in the top, WORKS!!
                docx_replace_regex(dxDoc, regex1, io.BytesIO(images[r]), r)
        
        ifConditions=[]
        findIfCondition(dxDoc)
        
        wFileName, wFileExt = os.path.splitext(args.output)
        wOutFile = wFileName + str(recs_cnt[0]) + wFileExt
        dxDoc.save(wOutFile)
        
        if mailServer:
            send_email(mailServer, cfg["Mailer"]['from_name'], cfg["Mailer"]['from_addr'], cfg["Mailer"]['subject'], mailMsg_html, mailMsg_text, rec_mas, wOutFile, args.attachments, args.images_attached)
        
        images={}
        regEx={}
        img_cnt = 0
        
        #make everything (except master query) as it was in configuration
        for i in range(1, len(cfg["Queries"])):
            mmm3_query[i] = list(cfg["Queries"])[i]
            
        rec_mas, images, regEx = QueryOneToDict(1, images, regEx)
        
except Exception as e:
    myLogging("Error", e)
    exc_type, exc_obj, exc_tb = sys.exc_info()
    fname = exc_tb.tb_frame.f_code.co_filename
    myLogging("Error", str(exc_type) + ", " + fname + " at line Number: " + str(exc_tb.tb_lineno))
    traceback.print_exc()
finally:
    # release cursor & connection
    if mas_cur:
        #same method for both Oracle and MySQL
        mas_cur.close()
        myLogging("Debug", "Cursor closed")
    if mas_connection:
        #same method for both Oracle and MySQL
        mas_connection.close()
        myLogging("Debug", "Connection closed")
    if log_file:
        log_file.close()
    if mailServer:
        mailServer.quit()
